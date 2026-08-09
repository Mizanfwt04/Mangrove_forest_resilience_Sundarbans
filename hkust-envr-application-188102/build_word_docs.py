#!/usr/bin/env python3
"""Convert application markdown files to Microsoft Word (.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

DIR = Path(__file__).parent
FILES = [
    ("cover_letter.md", "Cover_Letter_Rahman_HKUST_ENVR.docx"),
    ("research_statement.md", "Research_Statement_Rahman_HKUST_ENVR.docx"),
    ("teaching_statement.md", "Teaching_Statement_Rahman_HKUST_ENVR.docx"),
]


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_formatted_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def parse_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    set_default_style(doc)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|[-:\s|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue

        if in_table and table_rows:
            t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            t.style = "Table Grid"
            for r, row in enumerate(table_rows):
                for c, cell in enumerate(row):
                    t.rows[r].cells[c].text = cell
            table_rows = []
            in_table = False

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- [ ]"):
            add_formatted_paragraph(doc, "☐ " + stripped[5:])
        elif stripped.startswith("- "):
            add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
        elif stripped == "":
            pass
        else:
            add_formatted_paragraph(doc, stripped)

        i += 1

    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        t.style = "Table Grid"
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                t.rows[r].cells[c].text = cell

    doc.save(docx_path)
    print(f"Created {docx_path.name}")


def main() -> None:
    for md_name, docx_name in FILES:
        parse_markdown_to_docx(DIR / md_name, DIR / docx_name)


if __name__ == "__main__":
    main()
